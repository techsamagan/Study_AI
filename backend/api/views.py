from rest_framework import viewsets, status
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework_simplejwt.views import TokenObtainPairView
from django.contrib.auth import get_user_model
from django.db.models import Avg
from django.utils import timezone
import PyPDF2
import docx
import os

from .models import Document, Summary, Flashcard
from .serializers import (
    UserRegistrationSerializer, UserSerializer, CustomTokenObtainPairSerializer,
    DocumentSerializer, SummarySerializer, FlashcardSerializer, FlashcardCreateSerializer
)
from .services import OpenAIService
User = get_user_model()


class CustomTokenObtainPairView(TokenObtainPairView):
    """Custom token view that returns user data"""
    serializer_class = CustomTokenObtainPairSerializer

    def post(self, request, *args, **kwargs):
        response = super().post(request, *args, **kwargs)
        if response.status_code == 200:
            user = User.objects.get(email=request.data['email'])
            response.data['user'] = UserSerializer(user).data
        return response


@api_view(['POST'])
@permission_classes([AllowAny])
def register(request):
    """User registration endpoint"""
    serializer = UserRegistrationSerializer(data=request.data)
    if serializer.is_valid():
        user = serializer.save()
        return Response({
            'message': 'User registered successfully',
            'user': UserSerializer(user).data
        }, status=status.HTTP_201_CREATED)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET', 'PUT', 'PATCH'])
@permission_classes([IsAuthenticated])
def user_profile(request):
    """Get or update current user profile"""
    if request.method == 'GET':
        serializer = UserSerializer(request.user)
        return Response(serializer.data)
    elif request.method in ['PUT', 'PATCH']:
        serializer = UserSerializer(request.user, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class DocumentViewSet(viewsets.ModelViewSet):
    """ViewSet for document management"""
    serializer_class = DocumentSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Document.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        file = self.request.FILES.get('file')
        if not file:
            raise ValueError("No file provided")
        
        # Determine file type
        file_type = file.content_type or 'application/octet-stream'
        
        # Try to extract text and count pages
        pages = None
        try:
            if file_type == 'application/pdf':
                pdf_reader = PyPDF2.PdfReader(file)
                pages = len(pdf_reader.pages)
            elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                doc = docx.Document(file)
                # Estimate pages (rough calculation)
                pages = max(1, len(doc.paragraphs) // 20)
        except Exception:
            pass  # If extraction fails, pages remains None
        
        serializer.save(
            user=self.request.user,
            file_size=file.size,
            file_type=file_type,
            pages=pages
        )

    @action(detail=True, methods=['post'])
    def generate_summary(self, request, pk=None):
        """Generate summary for a document"""
        document = self.get_object()
    

        if document.user != request.user:
            return Response(
                {'error': 'Permission denied'},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            # Extract text from document
            text_content = self._extract_text_from_document(document)

            if not text_content:
                return Response(
                    {'error': 'Could not extract text from document'},
                    status=status.HTTP_400_BAD_REQUEST
                )


            # Generate summary using OpenAI
            openai_service = OpenAIService()

            summary_data = openai_service.generate_summary(text_content)

            # Create summary object
            summary = Summary.objects.create(
                user=request.user,
                document=document,
                full_summary=summary_data['full_summary'],
                key_points=summary_data.get('key_points', [])
            )

            serializer = SummarySerializer(summary)
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['post'])
    def generate_flashcards(self, request, pk=None):
        """Generate flashcards from a document"""
        document = self.get_object()
        
        if document.user != request.user:
            return Response(
                {'error': 'Permission denied'},
                status=status.HTTP_403_FORBIDDEN
            )

        num_cards = request.data.get('num_cards', 10)
        try:
            num_cards = int(num_cards)
            if num_cards < 1 or num_cards > 50:
                num_cards = 10
        except (ValueError, TypeError):
            num_cards = 10

        try:
            # Extract text from document
            text_content = self._extract_text_from_document(document)
            
            if not text_content:
                return Response(
                    {'error': 'Could not extract text from document'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Generate flashcards using OpenAI
            openai_service = OpenAIService()
            flashcards_data = openai_service.generate_flashcards(text_content, num_cards)

            # Create flashcard objects
            created_flashcards = []
            for card_data in flashcards_data:
                flashcard = Flashcard.objects.create(
                    user=request.user,
                    document=document,
                    question=card_data['question'],
                    answer=card_data['answer'],
                    category=card_data.get('category', 'General')
                )
                created_flashcards.append(flashcard)

            serializer = FlashcardSerializer(created_flashcards, many=True)
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _extract_text_from_document(self, document):
        """Extract text content from document file"""
        try:
            file_path = document.file.path
            file_type = document.file_type

            if file_type == 'application/pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ''
                    for page in pdf_reader.pages:
                        text += page.extract_text() + '\n'
                    return text

            elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                doc = docx.Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                return text

            elif file_type == 'text/plain':
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()

            else:
                # Try to read as text anyway
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        return file.read()
                except:
                    return None

        except Exception as e:
            print(f"Error extracting text: {str(e)}")
            return None


class SummaryViewSet(viewsets.ModelViewSet):
    """ViewSet for summary management"""
    serializer_class = SummarySerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Summary.objects.filter(user=self.request.user)

    @action(detail=True, methods=['post'])
    def generate_flashcards(self, request, pk=None):
        """Generate flashcards from a summary"""
        summary = self.get_object()
        
        if summary.user != request.user:
            return Response(
                {'error': 'Permission denied'},
                status=status.HTTP_403_FORBIDDEN
            )

        num_cards = request.data.get('num_cards', 10)
        try:
            num_cards = int(num_cards)
            if num_cards < 1 or num_cards > 50:
                num_cards = 10
        except (ValueError, TypeError):
            num_cards = 10

        try:
            # Use summary content to generate flashcards
            text_content = summary.full_summary
            
            # Generate flashcards using OpenAI
            openai_service = OpenAIService()
            flashcards_data = openai_service.generate_flashcards(text_content, num_cards)

            # Create flashcard objects
            created_flashcards = []
            for card_data in flashcards_data:
                flashcard = Flashcard.objects.create(
                    user=request.user,
                    document=summary.document,
                    summary=summary,
                    question=card_data['question'],
                    answer=card_data['answer'],
                    category=card_data.get('category', 'General')
                )
                created_flashcards.append(flashcard)

            serializer = FlashcardSerializer(created_flashcards, many=True)
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class FlashcardViewSet(viewsets.ModelViewSet):
    """ViewSet for flashcard management"""
    serializer_class = FlashcardSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        queryset = Flashcard.objects.filter(user=self.request.user)
        
        # Filter by category if provided
        category = self.request.query_params.get('category', None)
        if category:
            queryset = queryset.filter(category=category)
        
        # Filter by document if provided
        document_id = self.request.query_params.get('document', None)
        if document_id:
            queryset = queryset.filter(document_id=document_id)
        
        return queryset

    def get_serializer_class(self):
        if self.action == 'create':
            return FlashcardCreateSerializer
        return FlashcardSerializer

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    @action(detail=True, methods=['post'])
    def review(self, request, pk=None):
        """Mark flashcard as reviewed"""
        flashcard = self.get_object()
        
        if flashcard.user != request.user:
            return Response(
                {'error': 'Permission denied'},
                status=status.HTTP_403_FORBIDDEN
            )

        mastery_level = request.data.get('mastery_level', flashcard.mastery_level)
        try:
            mastery_level = int(mastery_level)
            if mastery_level < 0:
                mastery_level = 0
            elif mastery_level > 100:
                mastery_level = 100
        except (ValueError, TypeError):
            mastery_level = flashcard.mastery_level

        flashcard.last_reviewed = timezone.now()
        flashcard.review_count += 1
        flashcard.mastery_level = mastery_level
        flashcard.save()

        serializer = self.get_serializer(flashcard)
        return Response(serializer.data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def dashboard_stats(request):
    """Get dashboard statistics"""
    user = request.user
    
    mastery_result = Flashcard.objects.filter(user=user).aggregate(
        avg_mastery=Avg('mastery_level')
    )
    # The key is 'avg_mastery' when using aggregate(avg_mastery=Avg(...))
    mastery_value = mastery_result.get('avg_mastery') or 0
    
    stats = {
        'documents_count': Document.objects.filter(user=user).count(),
        'flashcards_count': Flashcard.objects.filter(user=user).count(),
        'summaries_count': Summary.objects.filter(user=user).count(),
        'study_time': '0h',  # Can be implemented with study session tracking
        'mastery': int(mastery_value) if mastery_value else 0,
    }
    
    return Response(stats)

